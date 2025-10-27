from __future__ import annotations

from abc import ABC, abstractmethod
from io import BytesIO
from typing import Any, Dict, Iterable, List, Tuple

from openpyxl import Workbook
from openpyxl.utils import get_column_letter
from openpyxl.styles import Font, PatternFill
from docx import Document
from docx.shared import Inches
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer

from .services.inventory import InventoryReportRow
from .csv_exporters import (
    SalesCSVExporter,
    CustomerCSVExporter,
    InventoryCSVExporter,
    AuditLogCSVExporter,
)
from .pdf_exporters import (
    SalesPDFExporter,
    CustomerPDFExporter,
    InventoryPDFExporter,
    AuditLogPDFExporter,
)


class BaseReportExporter(ABC):
    content_type: str
    file_extension: str

    @abstractmethod
    def export(self, report_data: Dict[str, Any]) -> bytes:  # pragma: no cover - interface
        raise NotImplementedError


class ExcelReportExporter(BaseReportExporter):
    content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    file_extension = 'xlsx'

    def export(self, report_data: Dict[str, Any]) -> bytes:
        workbook = Workbook()
        summary_sheet = workbook.active
        summary_sheet.title = 'Summary'

        summary_sheet.append(['Inventory Valuation Report'])
        summary_sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=2)
        summary_sheet['A1'].font = Font(size=14, bold=True)

        summary_sheet.append(['Generated At', report_data['generated_at'].strftime('%Y-%m-%d %H:%M:%S %Z')])
        if report_data.get('filters'):
            for key, value in report_data['filters'].items():
                summary_sheet.append([key.replace('_', ' ').title(), str(value)])
        summary_sheet.append([])

        summary_sheet.append(['Metric', 'Value'])
        for metric, value in report_data['summary'].items():
            label = metric.replace('_', ' ').title()
            summary_sheet.append([label, value])

        detail_sheet = workbook.create_sheet(title='Inventory Detail')
        detail_sheet.append(report_data['detail_headers'])

        for row in report_data['rows']:
            detail_sheet.append(row.as_list())

        self._auto_fit_columns([summary_sheet, detail_sheet])

        with BytesIO() as output:
            workbook.save(output)
            return output.getvalue()

    @staticmethod
    def _auto_fit_columns(sheets: Iterable[Any]) -> None:
        for sheet in sheets:
            for column_cells in sheet.columns:
                max_length = 0
                column = get_column_letter(column_cells[0].column)
                for cell in column_cells:
                    try:
                        cell_value = str(cell.value) if cell.value is not None else ''
                    except ValueError:
                        cell_value = ''
                    max_length = max(max_length, len(cell_value))
                adjusted_width = max_length + 2
                sheet.column_dimensions[column].width = min(adjusted_width, 50)


class WordReportExporter(BaseReportExporter):
    content_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    file_extension = 'docx'

    def export(self, report_data: Dict[str, Any]) -> bytes:
        document = Document()
        document.add_heading('Inventory Valuation Report', level=1)
        document.add_paragraph(f"Generated at: {report_data['generated_at'].strftime('%Y-%m-%d %H:%M:%S %Z')}")

        if report_data.get('filters'):
            document.add_heading('Filters', level=2)
            for key, value in report_data['filters'].items():
                document.add_paragraph(f"{key.replace('_', ' ').title()}: {value}", style='List Bullet')

        document.add_heading('Summary', level=2)
        for metric, value in report_data['summary'].items():
            document.add_paragraph(f"{metric.replace('_', ' ').title()}: {value}")

        document.add_heading('Inventory Detail', level=2)
        table = document.add_table(rows=1, cols=len(report_data['detail_headers']))
        header_cells = table.rows[0].cells
        for idx, header in enumerate(report_data['detail_headers']):
            header_cells[idx].text = header

        max_rows = 200  # prevent overly large documents
        for idx, row in enumerate(report_data['rows']):
            if idx >= max_rows:
                break
            cells = table.add_row().cells
            for cell_idx, value in enumerate(row.as_list()):
                cells[cell_idx].text = str(value)

        document.add_paragraph()
        document.add_paragraph('Note: Detail view trimmed to first 200 rows for readability.')

        with BytesIO() as output:
            document.save(output)
            return output.getvalue()


class PDFReportExporter(BaseReportExporter):
    content_type = 'application/pdf'
    file_extension = 'pdf'

    def export(self, report_data: Dict[str, Any]) -> bytes:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()

        story: List[Any] = []
        story.append(Paragraph('Inventory Valuation Report', styles['Title']))
        story.append(Paragraph(f"Generated at: {report_data['generated_at'].strftime('%Y-%m-%d %H:%M:%S %Z')}", styles['Normal']))
        story.append(Spacer(1, 12))

        if report_data.get('filters'):
            story.append(Paragraph('Filters', styles['Heading2']))
            for key, value in report_data['filters'].items():
                story.append(Paragraph(f"• {key.replace('_', ' ').title()}: {value}", styles['Normal']))
            story.append(Spacer(1, 12))

        summary_table_data = [['Metric', 'Value']]
        for metric, value in report_data['summary'].items():
            summary_table_data.append([metric.replace('_', ' ').title(), str(value)])
        summary_table = Table(summary_table_data, hAlign='LEFT')
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
            ('TOPPADDING', (0, 0), (-1, 0), 6),
            ('INNERGRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ('BOX', (0, 0), (-1, -1), 0.25, colors.grey),
        ]))
        story.append(summary_table)
        story.append(Spacer(1, 18))

        detail_table_data = [report_data['detail_headers']]
        max_rows = 60
        for idx, row in enumerate(report_data['rows']):
            if idx >= max_rows:
                break
            detail_table_data.append([str(value) for value in row.as_list()])
        detail_table = Table(detail_table_data, repeatRows=1)
        detail_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.darkblue),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 9),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
            ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
        ]))
        story.append(detail_table)
        if len(report_data['rows']) > max_rows:
            story.append(Spacer(1, 12))
            story.append(Paragraph('Detail view trimmed to first 60 rows for PDF output.', styles['Italic']))

        doc.build(story)
        buffer.seek(0)
        return buffer.read()


class SalesExcelExporter(BaseReportExporter):
    """Excel exporter specifically for sales data with multiple sheets"""
    content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    file_extension = 'xlsx'
    
    def export(self, report_data: Dict[str, Any]) -> bytes:
        workbook = Workbook()
        
        # Summary Sheet
        summary_sheet = workbook.active
        summary_sheet.title = 'Summary'
        
        # Header
        summary_sheet['A1'] = 'Sales Export Report'
        summary_sheet['A1'].font = Font(size=16, bold=True)
        summary_sheet.merge_cells('A1:B1')
        
        summary_sheet['A2'] = 'Generated At'
        summary_sheet['B2'] = report_data['generated_at'].strftime('%Y-%m-%d %H:%M:%S')
        
        row = 4
        summary_sheet[f'A{row}'] = 'Summary Metrics'
        summary_sheet[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        
        # Summary metrics
        summary_mapping = {
            'total_sales': 'Total Sales Count',
            'total_revenue': 'Total Revenue',
            'net_sales': 'Net Sales (excl. tax & discounts)',
            'total_tax': 'Total Tax Collected',
            'total_discounts': 'Total Discounts Given',
            'total_cogs': 'Total Cost of Goods Sold',
            'total_profit': 'Total Gross Profit',
            'profit_margin_percent': 'Profit Margin %',
            'amount_paid': 'Total Amount Paid',
            'amount_refunded': 'Total Amount Refunded',
            'outstanding_balance': 'Outstanding Balance',
        }
        
        for key, label in summary_mapping.items():
            value = report_data['summary'].get(key, '')
            summary_sheet[f'A{row}'] = label
            summary_sheet[f'B{row}'] = str(value)
            summary_sheet[f'A{row}'].font = Font(bold=True)
            row += 1
        
        # Sales Detail Sheet
        detail_sheet = workbook.create_sheet(title='Sales Detail')
        
        # Headers
        headers = [
            'Receipt Number', 'Date', 'Time', 'Storefront', 'Cashier',
            'Customer Name', 'Customer Type', 'Sale Type', 'Status',
            'Subtotal', 'Discount', 'Tax', 'Total', 
            'Amount Paid', 'Amount Refunded', 'Amount Due',
            'Payment Type', 'Notes'
        ]
        
        for col_num, header in enumerate(headers, 1):
            cell = detail_sheet.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='CCCCCC', end_color='CCCCCC', fill_type='solid')
        
        row = 2
        for sale in report_data['sales']:
            detail_sheet.cell(row=row, column=1, value=sale.get('receipt_number', ''))
            detail_sheet.cell(row=row, column=2, value=sale.get('date', ''))
            detail_sheet.cell(row=row, column=3, value=sale.get('time', ''))
            detail_sheet.cell(row=row, column=4, value=sale.get('storefront', ''))
            detail_sheet.cell(row=row, column=5, value=sale.get('cashier', ''))
            detail_sheet.cell(row=row, column=6, value=sale.get('customer_name', ''))
            detail_sheet.cell(row=row, column=7, value=sale.get('customer_type', ''))
            detail_sheet.cell(row=row, column=8, value=sale.get('sale_type', ''))
            detail_sheet.cell(row=row, column=9, value=sale.get('status', ''))
            detail_sheet.cell(row=row, column=10, value=sale.get('subtotal', ''))
            detail_sheet.cell(row=row, column=11, value=sale.get('discount', ''))
            detail_sheet.cell(row=row, column=12, value=sale.get('tax', ''))
            detail_sheet.cell(row=row, column=13, value=sale.get('total', ''))
            detail_sheet.cell(row=row, column=14, value=sale.get('amount_paid', ''))
            detail_sheet.cell(row=row, column=15, value=sale.get('amount_refunded', ''))
            detail_sheet.cell(row=row, column=16, value=sale.get('amount_due', ''))
            detail_sheet.cell(row=row, column=17, value=sale.get('payment_type', ''))
            detail_sheet.cell(row=row, column=18, value=sale.get('notes', ''))
            row += 1
        
        # Line Items Sheet
        items_sheet = workbook.create_sheet(title='Line Items')
        
        item_headers = [
            'Receipt Number', 'Product Name', 'SKU', 'Category',
            'Quantity', 'Unit Price', 'Total Price', 'COGS', 'Profit', 'Margin %'
        ]
        
        for col_num, header in enumerate(item_headers, 1):
            cell = items_sheet.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='CCCCCC', end_color='CCCCCC', fill_type='solid')
        
        row = 2
        for sale in report_data['sales']:
            for item in sale.get('items', []):
                items_sheet.cell(row=row, column=1, value=sale['receipt_number'])
                items_sheet.cell(row=row, column=2, value=item['product_name'])
                items_sheet.cell(row=row, column=3, value=item['sku'])
                items_sheet.cell(row=row, column=4, value=item['category'])
                items_sheet.cell(row=row, column=5, value=item['quantity'])
                items_sheet.cell(row=row, column=6, value=item['unit_price'])
                items_sheet.cell(row=row, column=7, value=item['total_price'])
                items_sheet.cell(row=row, column=8, value=item['cogs'])
                items_sheet.cell(row=row, column=9, value=item['profit'])
                items_sheet.cell(row=row, column=10, value=item['margin_percent'])
                row += 1
        
        # Auto-size columns
        for sheet in [summary_sheet, detail_sheet, items_sheet]:
            for column_cells in sheet.columns:
                max_length = 0
                column = get_column_letter(column_cells[0].column)
                for cell in column_cells:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                sheet.column_dimensions[column].width = adjusted_width
        
        # Save to bytes
        with BytesIO() as output:
            workbook.save(output)
            return output.getvalue()


class CustomerExcelExporter(BaseReportExporter):
    """Excel exporter for customer data with credit aging"""
    content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    file_extension = 'xlsx'
    
    def export(self, report_data: Dict[str, Any]) -> bytes:
        workbook = Workbook()
        
        # Summary Sheet
        summary_sheet = workbook.active
        summary_sheet.title = 'Summary'
        
        # Header
        summary_sheet['A1'] = 'Customer Export Report'
        summary_sheet['A1'].font = Font(size=16, bold=True)
        summary_sheet.merge_cells('A1:B1')
        
        summary_sheet['A2'] = 'Generated At'
        summary_sheet['B2'] = report_data['generated_at'].strftime('%Y-%m-%d %H:%M:%S')
        
        row = 4
        summary_sheet[f'A{row}'] = 'Customer Statistics'
        summary_sheet[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        
        # Summary metrics
        summary_mapping = {
            'total_customers': 'Total Customers',
            'retail_customers': 'Retail Customers',
            'wholesale_customers': 'Wholesale Customers',
            'active_customers': 'Active Customers',
            'blocked_customers': 'Blocked Customers',
            'total_credit_limit': 'Total Credit Limit',
            'total_outstanding_balance': 'Total Outstanding Balance',
            'total_available_credit': 'Total Available Credit',
        }
        
        for key, label in summary_mapping.items():
            value = report_data['summary'].get(key, '')
            summary_sheet[f'A{row}'] = label
            summary_sheet[f'B{row}'] = str(value)
            summary_sheet[f'A{row}'].font = Font(bold=True)
            row += 1
        
        # Aging Summary
        row += 1
        summary_sheet[f'A{row}'] = 'Aging Analysis Summary'
        summary_sheet[f'A{row}'].font = Font(bold=True, size=12)
        row += 1
        
        aging_mapping = {
            'aging_current': 'Current (0-30 days)',
            'aging_31_60': '31-60 days',
            'aging_61_90': '61-90 days',
            'aging_over_90': 'Over 90 days',
            'total_overdue': 'Total Overdue',
        }
        
        for key, label in aging_mapping.items():
            value = report_data['summary'].get(key, '')
            summary_sheet[f'A{row}'] = label
            summary_sheet[f'B{row}'] = str(value)
            summary_sheet[f'A{row}'].font = Font(bold=True)
            row += 1
        
        # Customer Detail Sheet
        detail_sheet = workbook.create_sheet(title='Customer Details')
        
        # Headers
        headers = [
            'Customer ID', 'Name', 'Email', 'Phone', 'Address',
            'Customer Type', 'Contact Person',
            'Credit Limit', 'Outstanding Balance', 'Available Credit',
            'Credit Terms (days)', 'Credit Blocked',
            'Total Sales Count', 'Total Sales Amount', 'Average Sale',
            'Last Sale Date', 'First Sale Date',
            'Active', 'Created Date', 'Created By'
        ]
        
        for col_num, header in enumerate(headers, 1):
            cell = detail_sheet.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='CCCCCC', end_color='CCCCCC', fill_type='solid')
        
        row = 2
        for customer in report_data['customers']:
            detail_sheet.cell(row=row, column=1, value=customer.get('customer_id', ''))
            detail_sheet.cell(row=row, column=2, value=customer.get('name', ''))
            detail_sheet.cell(row=row, column=3, value=customer.get('email', ''))
            detail_sheet.cell(row=row, column=4, value=customer.get('phone', ''))
            detail_sheet.cell(row=row, column=5, value=customer.get('address', ''))
            detail_sheet.cell(row=row, column=6, value=customer.get('customer_type', ''))
            detail_sheet.cell(row=row, column=7, value=customer.get('contact_person', ''))
            detail_sheet.cell(row=row, column=8, value=customer.get('credit_limit', ''))
            detail_sheet.cell(row=row, column=9, value=customer.get('outstanding_balance', ''))
            detail_sheet.cell(row=row, column=10, value=customer.get('available_credit', ''))
            detail_sheet.cell(row=row, column=11, value=customer.get('credit_terms_days', ''))
            detail_sheet.cell(row=row, column=12, value=customer.get('credit_blocked', ''))
            detail_sheet.cell(row=row, column=13, value=customer.get('total_sales_count', ''))
            detail_sheet.cell(row=row, column=14, value=customer.get('total_sales_amount', ''))
            detail_sheet.cell(row=row, column=15, value=customer.get('average_sale_amount', ''))
            detail_sheet.cell(row=row, column=16, value=customer.get('last_sale_date', ''))
            detail_sheet.cell(row=row, column=17, value=customer.get('first_sale_date', ''))
            detail_sheet.cell(row=row, column=18, value=customer.get('is_active', ''))
            detail_sheet.cell(row=row, column=19, value=customer.get('created_at', ''))
            detail_sheet.cell(row=row, column=20, value=customer.get('created_by', ''))
            row += 1
        
        # Aging Report Sheet
        aging_sheet = workbook.create_sheet(title='Credit Aging')
        
        aging_headers = [
            'Customer Name', 'Customer Type', 'Credit Limit', 'Outstanding Balance',
            'Current (0-30)', '31-60 Days', '61-90 Days', 'Over 90 Days',
            'Total Overdue', 'Oldest Invoice (days)', 'Credit Blocked'
        ]
        
        for col_num, header in enumerate(aging_headers, 1):
            cell = aging_sheet.cell(row=1, column=col_num)
            cell.value = header
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='FFE6E6', end_color='FFE6E6', fill_type='solid')
        
        row = 2
        for customer in report_data['customers']:
            aging_sheet.cell(row=row, column=1, value=customer.get('name', ''))
            aging_sheet.cell(row=row, column=2, value=customer.get('customer_type', ''))
            aging_sheet.cell(row=row, column=3, value=customer.get('credit_limit', ''))
            aging_sheet.cell(row=row, column=4, value=customer.get('outstanding_balance', ''))
            aging_sheet.cell(row=row, column=5, value=customer.get('aging_current', ''))
            aging_sheet.cell(row=row, column=6, value=customer.get('aging_31_60', ''))
            aging_sheet.cell(row=row, column=7, value=customer.get('aging_61_90', ''))
            aging_sheet.cell(row=row, column=8, value=customer.get('aging_over_90', ''))
            aging_sheet.cell(row=row, column=9, value=customer.get('total_overdue', ''))
            aging_sheet.cell(row=row, column=10, value=customer.get('oldest_invoice_days', ''))
            aging_sheet.cell(row=row, column=11, value=customer.get('credit_blocked', ''))
            row += 1
        
        # Credit Transactions Sheet (if available)
        if report_data.get('credit_transactions'):
            txn_sheet = workbook.create_sheet(title='Credit Transactions')
            
            txn_headers = [
                'Customer Name', 'Date', 'Transaction Type',
                'Amount', 'Balance Before', 'Balance After', 'Reference'
            ]
            
            for col_num, header in enumerate(txn_headers, 1):
                cell = txn_sheet.cell(row=1, column=col_num)
                cell.value = header
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color='E6F2FF', end_color='E6F2FF', fill_type='solid')
            
            row = 2
            for txn in report_data['credit_transactions']:
                txn_sheet.cell(row=row, column=1, value=txn.get('customer_name', ''))
                txn_sheet.cell(row=row, column=2, value=txn.get('date', ''))
                txn_sheet.cell(row=row, column=3, value=txn.get('transaction_type', ''))
                txn_sheet.cell(row=row, column=4, value=txn.get('amount', ''))
                txn_sheet.cell(row=row, column=5, value=txn.get('balance_before', ''))
                txn_sheet.cell(row=row, column=6, value=txn.get('balance_after', ''))
                txn_sheet.cell(row=row, column=7, value=txn.get('reference', ''))
                row += 1
        
        # Auto-size columns
        for sheet in workbook.worksheets:
            for column_cells in sheet.columns:
                max_length = 0
                column = get_column_letter(column_cells[0].column)
                for cell in column_cells:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                sheet.column_dimensions[column].width = adjusted_width
        
        # Save to bytes
        with BytesIO() as output:
            workbook.save(output)
            return output.getvalue()


class InventoryExcelExporter:
    """Excel exporter for inventory data"""
    
    content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    file_extension = 'xlsx'
    
    def export(self, data: Dict[str, Any]) -> bytes:
        """Export inventory data to Excel format"""
        workbook = Workbook()
        
        # Remove default sheet
        workbook.remove(workbook.active)
        
        # Sheet 1: Summary
        summary_sheet = workbook.create_sheet('Summary')
        
        # Header
        summary_sheet['A1'] = 'Inventory Export Summary'
        summary_sheet['A1'].font = Font(size=14, bold=True)
        summary_sheet.merge_cells('A1:B1')
        
        # Export metadata
        summary_sheet['A3'] = 'Export Date:'
        summary_sheet['B3'] = data['summary'].get('export_date', '')
        summary_sheet['A3'].font = Font(bold=True)
        
        # Summary statistics section
        summary_sheet['A5'] = 'Inventory Statistics'
        summary_sheet['A5'].font = Font(size=12, bold=True)
        summary_sheet['A5'].fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        summary_sheet['B5'].fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        
        row = 6
        metrics = [
            ('Total Unique Products', 'total_unique_products'),
            ('Total Quantity in Stock', 'total_quantity_in_stock'),
            ('Total Inventory Value', 'total_inventory_value'),
            ('Out of Stock Items', 'out_of_stock_items'),
            ('Low Stock Items', 'low_stock_items'),
            ('In Stock Items', 'in_stock_items'),
            ('Number of Storefronts', 'storefronts_count'),
        ]
        
        for label, key in metrics:
            summary_sheet.cell(row=row, column=1, value=label)
            summary_sheet.cell(row=row, column=1).font = Font(bold=True)
            summary_sheet.cell(row=row, column=2, value=data['summary'].get(key, 0))
            row += 1
        
        # Storefront breakdown
        row += 1
        summary_sheet.cell(row=row, column=1, value='Storefront Breakdown')
        summary_sheet.cell(row=row, column=1).font = Font(size=12, bold=True)
        summary_sheet.cell(row=row, column=1).fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        summary_sheet.cell(row=row, column=2).fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        summary_sheet.cell(row=row, column=3).fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        summary_sheet.cell(row=row, column=4).fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        row += 1
        
        # Storefront headers
        summary_sheet.cell(row=row, column=1, value='Storefront')
        summary_sheet.cell(row=row, column=2, value='Items')
        summary_sheet.cell(row=row, column=3, value='Quantity')
        summary_sheet.cell(row=row, column=4, value='Value')
        for col in range(1, 5):
            summary_sheet.cell(row=row, column=col).font = Font(bold=True)
        row += 1
        
        # Storefront data
        idx = 1
        while f'storefront_{idx}_name' in data['summary']:
            summary_sheet.cell(row=row, column=1, value=data['summary'][f'storefront_{idx}_name'])
            summary_sheet.cell(row=row, column=2, value=data['summary'][f'storefront_{idx}_items'])
            summary_sheet.cell(row=row, column=3, value=data['summary'][f'storefront_{idx}_quantity'])
            summary_sheet.cell(row=row, column=4, value=data['summary'][f'storefront_{idx}_value'])
            row += 1
            idx += 1
        
        # Sheet 2: Stock Items
        items_sheet = workbook.create_sheet('Stock Items')
        
        # Headers
        headers = [
            'Product ID', 'Product Name', 'SKU', 'Barcode', 'Storefront',
            'Quantity in Stock', 'Reorder Level', 'Unit of Measure', 'Stock Status',
            'Unit Cost', 'Selling Price', 'Total Value', 'Profit Margin', 'Margin %',
            'Last Updated', 'Created At', 'Created By'
        ]
        
        for col, header in enumerate(headers, 1):
            cell = items_sheet.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='D3D3D3', end_color='D3D3D3', fill_type='solid')
        
        # Data rows
        row = 2
        for item in data.get('stock_items', []):
            items_sheet.cell(row=row, column=1, value=item.get('product_id', ''))
            items_sheet.cell(row=row, column=2, value=item.get('product_name', ''))
            items_sheet.cell(row=row, column=3, value=item.get('sku', ''))
            items_sheet.cell(row=row, column=4, value=item.get('barcode', ''))
            items_sheet.cell(row=row, column=5, value=item.get('storefront', ''))
            items_sheet.cell(row=row, column=6, value=item.get('quantity_in_stock', 0))
            items_sheet.cell(row=row, column=7, value=item.get('reorder_level', 0))
            items_sheet.cell(row=row, column=8, value=item.get('unit_of_measure', ''))
            items_sheet.cell(row=row, column=9, value=item.get('stock_status', ''))
            items_sheet.cell(row=row, column=10, value=item.get('unit_cost', ''))
            items_sheet.cell(row=row, column=11, value=item.get('selling_price', ''))
            items_sheet.cell(row=row, column=12, value=item.get('total_value', ''))
            items_sheet.cell(row=row, column=13, value=item.get('profit_margin', ''))
            items_sheet.cell(row=row, column=14, value=item.get('margin_percentage', ''))
            items_sheet.cell(row=row, column=15, value=item.get('last_updated', ''))
            items_sheet.cell(row=row, column=16, value=item.get('created_at', ''))
            items_sheet.cell(row=row, column=17, value=item.get('created_by', ''))
            row += 1
        
        # Sheet 3: Stock Movements (if included)
        if data.get('stock_movements'):
            movements_sheet = workbook.create_sheet('Stock Movements')
            
            # Headers
            movement_headers = [
                'Date', 'Product Name', 'SKU', 'Storefront', 'Adjustment Type',
                'Quantity Before', 'Quantity Adjusted', 'Quantity After', 'Reason', 'Performed By'
            ]
            
            for col, header in enumerate(movement_headers, 1):
                cell = movements_sheet.cell(row=1, column=col, value=header)
                cell.font = Font(bold=True)
                cell.fill = PatternFill(start_color='CCEBFF', end_color='CCEBFF', fill_type='solid')
            
            # Data rows
            row = 2
            for movement in data['stock_movements']:
                movements_sheet.cell(row=row, column=1, value=movement.get('date', ''))
                movements_sheet.cell(row=row, column=2, value=movement.get('product_name', ''))
                movements_sheet.cell(row=row, column=3, value=movement.get('sku', ''))
                movements_sheet.cell(row=row, column=4, value=movement.get('storefront', ''))
                movements_sheet.cell(row=row, column=5, value=movement.get('adjustment_type', ''))
                movements_sheet.cell(row=row, column=6, value=movement.get('quantity_before', ''))
                movements_sheet.cell(row=row, column=7, value=movement.get('quantity_adjusted', ''))
                movements_sheet.cell(row=row, column=8, value=movement.get('quantity_after', ''))
                movements_sheet.cell(row=row, column=9, value=movement.get('reason', ''))
                movements_sheet.cell(row=row, column=10, value=movement.get('performed_by', ''))
                row += 1
        
        # Auto-size columns
        for sheet in workbook.worksheets:
            for column_cells in sheet.columns:
                max_length = 0
                column = get_column_letter(column_cells[0].column)
                for cell in column_cells:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                sheet.column_dimensions[column].width = adjusted_width
        
        # Save to bytes
        with BytesIO() as output:
            workbook.save(output)
            return output.getvalue()


class AuditLogExcelExporter:
    """Excel exporter for audit log data"""
    
    content_type = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    file_extension = 'xlsx'
    
    def export(self, data: Dict[str, Any]) -> bytes:
        """Export audit logs to Excel format"""
        workbook = Workbook()
        
        # Remove default sheet
        workbook.remove(workbook.active)
        
        # Sheet 1: Summary
        summary_sheet = workbook.create_sheet('Summary')
        
        # Header
        summary_sheet['A1'] = 'Audit Log Export Summary'
        summary_sheet['A1'].font = Font(size=14, bold=True)
        summary_sheet.merge_cells('A1:B1')
        
        # Export metadata
        summary_sheet['A3'] = 'Export Date:'
        summary_sheet['B3'] = data['summary'].get('export_date', '')
        summary_sheet['A3'].font = Font(bold=True)
        
        summary_sheet['A4'] = 'Date Range:'
        summary_sheet['B4'] = f"{data['summary'].get('date_range_start', '')} to {data['summary'].get('date_range_end', '')}"
        summary_sheet['A4'].font = Font(bold=True)
        
        # Summary statistics section
        summary_sheet['A6'] = 'Audit Statistics'
        summary_sheet['A6'].font = Font(size=12, bold=True)
        summary_sheet['A6'].fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        summary_sheet['B6'].fill = PatternFill(start_color='E0E0E0', end_color='E0E0E0', fill_type='solid')
        
        row = 7
        metrics = [
            ('Total Events', 'total_events'),
            ('First Event', 'first_event_time'),
            ('Last Event', 'last_event_time'),
            ('Unique Users', 'unique_users'),
            ('Unique Event Types', 'unique_event_types'),
        ]
        
        for label, key in metrics:
            summary_sheet.cell(row=row, column=1, value=label)
            summary_sheet.cell(row=row, column=1).font = Font(bold=True)
            summary_sheet.cell(row=row, column=2, value=data['summary'].get(key, 0))
            row += 1
        
        # Event type breakdown
        row += 1
        summary_sheet.cell(row=row, column=1, value='Top Event Types')
        summary_sheet.cell(row=row, column=1).font = Font(size=12, bold=True)
        summary_sheet.cell(row=row, column=1).fill = PatternFill(start_color='FFE6CC', end_color='FFE6CC', fill_type='solid')
        summary_sheet.cell(row=row, column=2).fill = PatternFill(start_color='FFE6CC', end_color='FFE6CC', fill_type='solid')
        row += 1
        
        # Event headers
        summary_sheet.cell(row=row, column=1, value='Event Type')
        summary_sheet.cell(row=row, column=2, value='Count')
        for col in range(1, 3):
            summary_sheet.cell(row=row, column=col).font = Font(bold=True)
        row += 1
        
        # Event data
        idx = 1
        while f'event_{idx}_type' in data['summary']:
            summary_sheet.cell(row=row, column=1, value=data['summary'][f'event_{idx}_type'])
            summary_sheet.cell(row=row, column=2, value=data['summary'][f'event_{idx}_count'])
            row += 1
            idx += 1
        
        # User activity breakdown
        row += 1
        summary_sheet.cell(row=row, column=1, value='Top Users')
        summary_sheet.cell(row=row, column=1).font = Font(size=12, bold=True)
        summary_sheet.cell(row=row, column=1).fill = PatternFill(start_color='CCE5FF', end_color='CCE5FF', fill_type='solid')
        summary_sheet.cell(row=row, column=2).fill = PatternFill(start_color='CCE5FF', end_color='CCE5FF', fill_type='solid')
        row += 1
        
        # User headers
        summary_sheet.cell(row=row, column=1, value='User')
        summary_sheet.cell(row=row, column=2, value='Activity Count')
        for col in range(1, 3):
            summary_sheet.cell(row=row, column=col).font = Font(bold=True)
        row += 1
        
        # User data
        idx = 1
        while f'user_{idx}_name' in data['summary']:
            summary_sheet.cell(row=row, column=1, value=data['summary'][f'user_{idx}_name'])
            summary_sheet.cell(row=row, column=2, value=data['summary'][f'user_{idx}_count'])
            row += 1
            idx += 1
        
        # Sheet 2: Audit Logs
        logs_sheet = workbook.create_sheet('Audit Logs')
        
        # Headers
        headers = [
            'Log ID', 'Timestamp', 'Event Type', 'Event Label', 'User Email', 'User Name',
            'Entity Type', 'Entity ID', 'Entity Name', 'Description', 'Event Details',
            'IP Address', 'User Agent'
        ]
        
        for col, header in enumerate(headers, 1):
            cell = logs_sheet.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.fill = PatternFill(start_color='4472C4', end_color='4472C4', fill_type='solid')
            cell.font = Font(bold=True, color='FFFFFF')
        
        # Data rows
        row = 2
        for log in data.get('audit_logs', []):
            logs_sheet.cell(row=row, column=1, value=log.get('log_id', ''))
            logs_sheet.cell(row=row, column=2, value=log.get('timestamp', ''))
            logs_sheet.cell(row=row, column=3, value=log.get('event_type', ''))
            logs_sheet.cell(row=row, column=4, value=log.get('event_label', ''))
            logs_sheet.cell(row=row, column=5, value=log.get('user_email', ''))
            logs_sheet.cell(row=row, column=6, value=log.get('user_name', ''))
            logs_sheet.cell(row=row, column=7, value=log.get('entity_type', ''))
            logs_sheet.cell(row=row, column=8, value=log.get('entity_id', ''))
            logs_sheet.cell(row=row, column=9, value=log.get('entity_name', ''))
            logs_sheet.cell(row=row, column=10, value=log.get('description', ''))
            logs_sheet.cell(row=row, column=11, value=log.get('event_details', ''))
            logs_sheet.cell(row=row, column=12, value=log.get('ip_address', ''))
            logs_sheet.cell(row=row, column=13, value=log.get('user_agent', ''))
            row += 1
        
        # Auto-size columns
        for sheet in workbook.worksheets:
            for column_cells in sheet.columns:
                max_length = 0
                column = get_column_letter(column_cells[0].column)
                for cell in column_cells:
                    try:
                        if cell.value:
                            max_length = max(max_length, len(str(cell.value)))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 60)  # Max 60 for audit logs
                sheet.column_dimensions[column].width = adjusted_width
        
        # Save to bytes
        with BytesIO() as output:
            workbook.save(output)
            return output.getvalue()


EXPORTER_MAP = {
    'excel': ExcelReportExporter,
    'docx': WordReportExporter,
    'pdf': PDFReportExporter,
    'sales_excel': SalesExcelExporter,
    'sales_csv': SalesCSVExporter,
    'sales_pdf': SalesPDFExporter,
    'customer_excel': CustomerExcelExporter,
    'customer_csv': CustomerCSVExporter,
    'customer_pdf': CustomerPDFExporter,
    'inventory_excel': InventoryExcelExporter,
    'inventory_csv': InventoryCSVExporter,
    'inventory_pdf': InventoryPDFExporter,
    'audit_excel': AuditLogExcelExporter,
    'audit_csv': AuditLogCSVExporter,
    'audit_pdf': AuditLogPDFExporter,
}
